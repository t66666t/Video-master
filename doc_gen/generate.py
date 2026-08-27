# -*- coding: utf-8 -*-
"""
从 doc_content.py 的结构化内容稿生成 Word(.docx) 与 PDF(.pdf)。
用法: python generate.py
"""
import re
import sys
import os

from doc_content import TITLE, SUBTITLE, BLOCKS

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "Fluent_Player_项目全解.docx")
PDF_PATH = os.path.join(OUT_DIR, "Fluent_Player_项目全解.pdf")

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def parse_inline(text):
    """把 '**粗体**' 标记解析为 (text, is_bold) 段列表。"""
    parts = []
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    return parts


# =====================================================================
# Word 生成
# =====================================================================
def make_docx():
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    CJK_FONT = "微软雅黑"
    MONO_FONT = "Consolas"
    ACCENT = RGBColor(0x1F, 0x6F, 0xB2)
    GRAY = RGBColor(0x60, 0x60, 0x60)

    doc = Document()

    # ---- 全局样式 ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.4

    def _set_cjk(run, name=CJK_FONT):
        run.font.name = name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), name)

    def add_heading_runs(text, size, bold=True, color=None, space_before=12, space_after=8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        for seg, is_bold in parse_inline(text):
            run = p.add_run(seg)
            run.bold = bold or is_bold
            run.font.size = Pt(size)
            _set_cjk(run)
            if color is not None:
                run.font.color.rgb = color
        return p

    def shade_paragraph(p, fill="F2F2F2"):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def add_para(text, indent=True, size=11):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(6)
        if indent:
            p.paragraph_format.first_line_indent = Pt(22)
        for seg, is_bold in parse_inline(text):
            run = p.add_run(seg)
            run.bold = is_bold
            run.font.size = Pt(size)
            _set_cjk(run)
        return p

    # ---- 封面 ----
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(26)
    _set_cjk(run)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY
    _set_cjk(run)
    doc.add_page_break()

    # ---- 目录 ----
    p = add_heading_runs("目录", 18, space_before=6)
    toc = doc.add_paragraph()
    run = toc.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "（在 Word 中右键此处 → 更新域，即可生成目录）"
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)
    doc.add_page_break()

    # ---- 正文 ----
    for block in BLOCKS:
        kind = block[0]
        if kind == "pagebreak":
            doc.add_page_break()
        elif kind == "h1":
            p = add_heading_runs(block[1], 18, color=ACCENT, space_before=6, space_after=10)
            # 给一级标题加边框底纹
            shade_paragraph(p, "DEEBF7")
        elif kind == "h2":
            add_heading_runs(block[1], 15, color=ACCENT, space_before=14, space_after=6)
        elif kind == "h3":
            add_heading_runs(block[1], 12.5, color=RGBColor(0x33, 0x33, 0x33), space_before=10, space_after=4)
        elif kind == "p":
            add_para(block[1])
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.4
            p.paragraph_format.space_after = Pt(3)
            for seg, is_bold in parse_inline(block[1]):
                run = p.add_run(seg)
                run.bold = is_bold
                run.font.size = Pt(11)
                _set_cjk(run)
        elif kind == "num":
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.4
            p.paragraph_format.space_after = Pt(3)
            for seg, is_bold in parse_inline(block[1]):
                run = p.add_run(seg)
                run.bold = is_bold
                run.font.size = Pt(11)
                _set_cjk(run)
        elif kind == "code":
            for line in block[1].rstrip("\n").split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.left_indent = Pt(12)
                shade_paragraph(p, "F5F5F5")
                run = p.add_run(line if line else " ")
                run.font.name = MONO_FONT
                run.font.size = Pt(9.5)
                _set_cjk(run, MONO_FONT)
                run.font.color.rgb = RGBColor(0x24, 0x24, 0x24)
            # 代码块后加一点空隙
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
        elif kind == "note":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.4
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.right_indent = Pt(10)
            shade_paragraph(p, "FFF4E5")
            lead = p.add_run("提示：")
            lead.bold = True
            lead.font.color.rgb = RGBColor(0xB0, 0x59, 0x00)
            lead.font.size = Pt(10.5)
            _set_cjk(lead)
            for seg, is_bold in parse_inline(block[1]):
                run = p.add_run(seg)
                run.bold = is_bold
                run.font.size = Pt(10.5)
                _set_cjk(run)
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.4
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.right_indent = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, "F0F4F8")
            for seg, is_bold in parse_inline(block[1]):
                run = p.add_run(seg)
                run.bold = is_bold
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                _set_cjk(run)
        elif kind == "table":
            headers, rows = block[1], block[2]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(headers):
                cell = table.cell(0, j)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(10)
                _set_cjk(run)
                # 表头底纹
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "DEEBF7")
                tcPr.append(shd)
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    cell = table.cell(i + 1, j)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    for seg, is_bold in parse_inline(str(val)):
                        run = p.add_run(seg)
                        run.bold = is_bold
                        run.font.size = Pt(9.5)
                        _set_cjk(run)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.save(DOCX_PATH)
    return DOCX_PATH


# =====================================================================
# PDF 生成
# =====================================================================
def make_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Preformatted, Table, TableStyle,
        PageBreak, KeepTogether, ListFlowable, ListItem,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdfmetrics.registerFont(TTFont("Msyh", "C:/Windows/Fonts/msyh.ttc", subfontIndex=0))
    pdfmetrics.registerFont(TTFont("Msyh-Bold", "C:/Windows/Fonts/msyhbd.ttc", subfontIndex=0))
    pdfmetrics.registerFontFamily("Msyh", normal="Msyh", bold="Msyh-Bold", italic="Msyh", boldItalic="Msyh-Bold")

    ACCENT = colors.HexColor("#1F6FB2")
    DARK = colors.HexColor("#242424")
    GRAY = colors.HexColor("#606060")

    def s(name, **kw):
        base = dict(
            fontName="Msyh", fontSize=10.5, leading=16,
            textColor=DARK, spaceAfter=6, wordWrap="CJK",
        )
        base.update(kw)
        return ParagraphStyle(name, **base)

    styles = {
        "h1": s("h1", fontSize=17, leading=24, fontName="Msyh-Bold",
                textColor=ACCENT, spaceBefore=6, spaceAfter=10),
        "h2": s("h2", fontSize=14, leading=20, fontName="Msyh-Bold",
                textColor=ACCENT, spaceBefore=14, spaceAfter=6),
        "h3": s("h3", fontSize=12, leading=17, fontName="Msyh-Bold",
                textColor=colors.HexColor("#333333"), spaceBefore=10, spaceAfter=4),
        "p": s("p", firstLineIndent=22),
        "bullet": s("bullet", leftIndent=20, bulletIndent=8, spaceAfter=3),
        "num": s("num", leftIndent=20, bulletIndent=8, spaceAfter=3),
        "note": s("note", fontSize=10, leading=15, leftIndent=8, rightIndent=8,
                  textColor=colors.HexColor("#6B4A00")),
        "quote": s("quote", fontSize=10, leading=15, leftIndent=20, rightIndent=12,
                   textColor=colors.HexColor("#404040")),
        "code": s("code", fontSize=8.6, leading=12.5, fontName="Msyh",
                  textColor=colors.HexColor("#242424"), leftIndent=10, spaceAfter=0),
    }

    def esc(text):
        return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def rich(text):
        """把 **bold** 转成 <b>。</b>"""
        out = []
        pos = 0
        for m in BOLD_RE.finditer(text):
            if m.start() > pos:
                out.append(esc(text[pos:m.start()]))
            out.append("<b>%s</b>" % esc(m.group(1)))
            pos = m.end()
        if pos < len(text):
            out.append(esc(text[pos:]))
        return "".join(out)

    story = []

    # ---- 封面 ----
    for _ in range(7):
        story.append(Spacer(1, 18))
    story.append(Paragraph(esc(TITLE), s("cover_title", fontSize=24, leading=34,
                                        fontName="Msyh-Bold", alignment=1)))
    story.append(Spacer(1, 14))
    story.append(Paragraph(esc(SUBTITLE), s("cover_sub", fontSize=13, leading=20,
                                            textColor=GRAY, alignment=1)))
    story.append(PageBreak())

    # ---- 目录占位 ----
    story.append(Paragraph("目录", styles["h1"]))
    story.append(Paragraph("（PDF 预览版目录为静态文字；Word 版可在 Word 中右键“更新域”生成可点击目录）",
                           s("toc_note", fontSize=9, textColor=GRAY)))
    toc_items = []
    for block in BLOCKS:
        if block[0] == "h1":
            toc_items.append((1, block[1]))
        elif block[0] == "h2":
            toc_items.append((2, block[1]))
    for level, txt in toc_items:
        indent = "　　" if level == 2 else ""
        story.append(Paragraph(esc(indent + txt), s("toc", fontSize=10.5,
                                                    leftIndent=(level - 1) * 14,
                                                    spaceAfter=3)))
    story.append(PageBreak())

    # ---- 正文 ----
    for block in BLOCKS:
        kind = block[0]
        if kind == "pagebreak":
            story.append(PageBreak())
        elif kind in ("h1", "h2", "h3"):
            story.append(KeepTogether([Paragraph(rich(block[1]), styles[kind])]))
        elif kind == "p":
            story.append(Paragraph(rich(block[1]), styles["p"]))
        elif kind == "bullet":
            story.append(ListFlowable(
                [ListItem(Paragraph(rich(block[1]), styles["bullet"]))],
                bulletType="bullet", start="•",
                leftIndent=18, bulletFontName="Msyh",
            ))
        elif kind == "num":
            story.append(ListFlowable(
                [ListItem(Paragraph(rich(block[1]), styles["num"]))],
                bulletType="1", start=block[2] if len(block) > 2 else None,
                leftIndent=18, bulletFontName="Msyh",
            ))
        elif kind == "code":
            lines = block[1].rstrip("\n").split("\n")
            # 用带背景的单格 Table 包裹 Preformatted
            code_paras = [Paragraph(esc(line) if line else "&nbsp;", styles["code"])
                          for line in lines]
            inner = Table([[code_paras[0]]], colWidths=[None])
            # 更简单：直接 Preformatted
            pre = Preformatted(block[1].rstrip("\n"), styles["code"])
            wrapper = Table([[pre]], colWidths=[None])
            wrapper.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5F5F5")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(Spacer(1, 3))
            story.append(wrapper)
            story.append(Spacer(1, 5))
        elif kind == "note":
            note_p = Paragraph("<b>提示：</b>" + rich(block[1]), styles["note"])
            t = Table([[note_p]], colWidths=[None])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF4E5")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#F0C070")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 2))
            story.append(t)
            story.append(Spacer(1, 6))
        elif kind == "quote":
            quote_p = Paragraph(rich(block[1]), styles["quote"])
            t = Table([[quote_p]], colWidths=[None])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F0F4F8")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#C9D6E3")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 2))
            story.append(t)
            story.append(Spacer(1, 6))
        elif kind == "table":
            headers, rows = block[1], block[2]
            data = [[Paragraph("<b>%s</b>" % esc(str(h)), styles["p"]) for h in headers]]
            data += [[Paragraph(rich(str(v)), s("cell", fontSize=9, leading=13, spaceAfter=0))
                      for v in row] for row in rows]
            col_w = [22 * mm, 40 * mm, None]
            avail = 175 * mm
            col_w = [avail * (0.16 if i == 0 else (0.30 if i == 1 else 0.54)) for i in range(len(headers))]
            t = Table(data, colWidths=col_w, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C0C0C0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DEEBF7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(Spacer(1, 2))
            story.append(t)
            story.append(Spacer(1, 6))

    def _footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Msyh", 8.5)
        canvas.setFillColor(GRAY)
        canvas.drawCentredString(A4[0] / 2, 10 * mm, "%d" % doc_.page)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        PDF_PATH, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=16 * mm,
        title=TITLE, author="hy3",
    )
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return PDF_PATH


if __name__ == "__main__":
    p1 = make_docx()
    p2 = make_pdf()
    print("DOCX:", p1)
    print("PDF :", p2)
